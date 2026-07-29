from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INPUT = Path(r"D:\Documents\Luận văn tốt nghiệp\Báo cáo luận văn 4 - Copy.docx")
OUTPUT = Path(r"D:\VSC\TTTN_E-Learning-Platform\Báo cáo luận văn 4 - Copy_header_footer_2026.docx")
TITLE = "XÂY DỰNG WEBSITE NỀN TẢNG HỌC LẬP TRÌNH"


def clear_container(container):
    for paragraph in list(container.paragraphs):
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
    for table in list(container.tables):
        container._element.remove(table._element)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, cached, end])


def set_borderless(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_header(section, text):
    section.header.is_linked_to_previous = False
    clear_container(section.header)
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.upper())
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.font.italic = True


def set_footer(section):
    section.footer.is_linked_to_previous = False
    clear_container(section.footer)
    table = section.footer.add_table(rows=1, cols=2, width=Cm(16))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13)
    table.columns[1].width = Cm(3)
    for cell in table.rows[0].cells:
        set_borderless(cell)
    left, right = table.rows[0].cells
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = left.paragraphs[0].add_run(TITLE)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.font.italic = True
    add_page_field(right.paragraphs[0])


def add_section_before(doc, paragraph_index, source_section_index):
    paragraph = doc.paragraphs[paragraph_index]
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.sectPr is not None:
        return
    sect_pr = deepcopy(doc.sections[source_section_index]._sectPr)
    # Do not reuse the previous section's header/footer parts. New parts are
    # created below so Chapter 3 and Chapter 4 can have different headers.
    for tag in ("w:headerReference", "w:footerReference"):
        for reference in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(reference)
    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    p_pr.append(sect_pr)


def clear_page_number_start(section):
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is not None:
        section._sectPr.remove(pg_num)


def set_page_number_start(section, value=1):
    clear_page_number_start(section)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(value))
    section._sectPr.append(pg_num)


def main():
    doc = Document(str(INPUT))

    # The original file has Chapter 3 and Chapter 4 in one section.
    # Add a section break before Chapter 4 so each chapter can have its own header.
    add_section_before(doc, 605, 6)

    for index, section in enumerate(doc.sections):
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        clear_container(section.header)
        clear_container(section.first_page_header)
        clear_container(section.footer)
        clear_container(section.first_page_footer)

        if index < 4:
            # Cover, acknowledgements and table-of-contents pages have no page furniture.
            continue

        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].text = ""
        section.first_page_footer.paragraphs[0].text = ""
        set_footer(section)

    headers = {
        4: "Chương 1. GIỚI THIỆU",
        5: "Chương 2. PHƯƠNG PHÁP THỰC HIỆN",
        6: "Chương 3. THIẾT KẾ",
        7: "Chương 4. THỬ NGHIỆM VÀ ĐÁNH GIÁ",
        8: "Chương 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        9: "TÀI LIỆU THAM KHẢO",
    }
    for index, text in headers.items():
        if index < len(doc.sections):
            set_header(doc.sections[index], text)

    if len(doc.sections) > 4:
        set_page_number_start(doc.sections[4], 1)
        for section in doc.sections[5:]:
            clear_page_number_start(section)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(str(OUTPUT))
    print("header and footer updated")


if __name__ == "__main__":
    main()
